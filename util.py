from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def save_to_word(results, date):
    doc = Document()

    # Tiêu đề chính
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(f"BẢN TIN NGÀY {date.upper()}\n")
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_paragraph.paragraph_format.space_after = Pt(5)

    # Chia danh sách tin tức
    others_news = [article for article in results if article["scale"] == "others"]
    vn_news = [article for article in results if article["scale"] == "VN"]

    # Chia tin trong nước theo danh mục
    chinh_tri_news = [article for article in vn_news if article["category"] == "Chính trị"]
    kinh_te_news = [article for article in vn_news if article["category"] == "Kinh tế"]
    tai_chinh_news = [article for article in vn_news if article["category"] == "Tài chính"]

    # Biến đếm thứ tự bài báo xuyên suốt
    article_index = 1

    # Xử lý tin quốc tế
    section_title = doc.add_paragraph()
    section_run = section_title.add_run("TIN QUỐC TẾ")
    section_run.bold = True
    section_run.font.size = Pt(14)
    section_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    section_title.paragraph_format.space_after = Pt(5)

    if others_news:
        for article in others_news:
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(f"{article_index}. {article['title']}")
            title_run.bold = True
            title_run.underline = True
            title_run.font.size = Pt(12)
            title_paragraph.paragraph_format.space_after = Pt(5)

            for line in article["content"].split("\n"):
                content_paragraph = doc.add_paragraph()
                content_run = content_paragraph.add_run(line)
                content_run.font.size = Pt(11)
                content_paragraph.paragraph_format.space_after = Pt(5)
            
            article_index += 1  
    else:
        no_news = doc.add_paragraph("(Không có tin quốc tế)")
        no_news.runs[0].italic = True
        no_news.paragraph_format.space_after = Pt(5)

    # Xử lý tin trong nước
    section_title = doc.add_paragraph()
    section_run = section_title.add_run("TIN TRONG NƯỚC")
    section_run.bold = True
    section_run.font.size = Pt(14)
    section_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    section_title.paragraph_format.space_after = Pt(5)

    # Hàm in từng danh mục với thứ tự liên tục
    def print_news_section(doc, section_name, news_list, start_index):
        section_paragraph = doc.add_paragraph()
        section_run = section_paragraph.add_run(section_name.upper())
        section_run.bold = True
        section_run.font.size = Pt(12)
        section_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        section_paragraph.paragraph_format.space_after = Pt(5)

        current_index = start_index
        if news_list:
            for article in news_list:
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(f"{current_index}. {article['title']}")
                title_run.bold = True
                title_run.underline = True
                title_run.font.size = Pt(13)
                title_paragraph.paragraph_format.space_after = Pt(0) 

                lines = article["content"].split("\n")
                for i, line in enumerate(lines):
                    content_paragraph = doc.add_paragraph()
                    content_run = content_paragraph.add_run(line)
                    content_run.font.size = Pt(12)
                    
                    if i == len(lines) - 1:
                        content_paragraph.paragraph_format.space_after = Pt(5)
                    else:
                        content_paragraph.paragraph_format.space_after = Pt(0)
                
                current_index += 1  
        else:
            no_news = doc.add_paragraph("(Không có tin thuộc mục này)")
            no_news.runs[0].italic = True
            no_news.paragraph_format.space_after = Pt(5)

        return current_index  

    # Gọi hàm in từng danh mục với thứ tự liên tục
    article_index = print_news_section(doc, "Chính trị", chinh_tri_news, article_index)
    article_index = print_news_section(doc, "Kinh tế", kinh_te_news, article_index)
    article_index = print_news_section(doc, "Tài chính", tai_chinh_news, article_index)

    # Lưu vào BytesIO thay vì ổ cứng
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)  # Đưa con trỏ về đầu file

    return word_buffer
